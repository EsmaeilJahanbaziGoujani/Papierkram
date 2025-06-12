import csv
import pandas as pd
from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog
import os

def create_weekly_report(data, output_path):
    doc = Document()

    # Header information
    doc.add_heading("Weekly Report", level=1)

    # Create a table for each date
    for date, details in data.items():
        doc.add_heading(f"Date: {date}", level=2)  # Add date as a subheading

        table = doc.add_table(rows=1, cols=2)  # Simplified table for key-value pairs
        table.style = 'Table Grid'

        # Add data to the table
        row_cells = table.add_row().cells
        row_cells[0].text = "Beginn"
        row_cells[1].text = str(details.get("Beginn", "")) if details.get("Beginn") else ""

        row_cells = table.add_row().cells
        row_cells[0].text = "Ende"
        row_cells[1].text = str(details.get("Ende", "")) if details.get("Ende") else ""

        row_cells = table.add_row().cells
        row_cells[0].text = "Dauer"
        row_cells[1].text = str(details.get("Dauer", "")) if details.get("Dauer") else ""
        # Adding comments which can be extended to more fields if needed

        row_cells = table.add_row().cells
        row_cells[0].text = "Kommentar"
        row_cells[1].text = str(details.get("Kommentar", "")) if details.get("Kommentar") else ""

    doc.save(output_path)

def browse_file():
    try:
        filename = filedialog.askopenfilename(initialdir=".",
                                                title="Select CSV File",
                                                filetypes=[("CSV files", "*.csv"), ("All files", "*.*")])
        if filename:  # Check if a file was selected
            file_path_entry.delete(0, tk.END)
            file_path_entry.insert(0, filename)
    except Exception as e:
        status_label.config(text=f"Error during file selection: {str(e)}")

def convert_to_word():
    file_path = file_path_entry.get()
    if not file_path:
        status_label.config(text="Please select a CSV file.")
        return

    try:
        data = pd.read_csv(file_path, delimiter=';') # Reading the sample csv file

        # Creating a dictionary for the weekly report
        weekly_data = {}

        for index, row in data.iterrows():
            date = row['Datum']

            # Check for NaNs and handle them, convert value to string
            weekly_data[date] = {
                'Beginn': str(row['Beginn']) if not pd.isna(row['Beginn']) else '',
                'Ende': str(row['Ende']) if not pd.isna(row['Ende']) else '',
                'Dauer': str(row['Dauer (hh:mm)']) if not pd.isna(row['Dauer (hh:mm)']) else '',
                'Dauer_abrechenbar': str(row['Dauer (abrechenbar) (hh:mm)']) if not pd.isna(row['Dauer (abrechenbar) (hh:mm)']) else '',
                'Nicht_Abrechenbar': str(row['Nicht-Abrechenbar']) if not pd.isna(row['Nicht-Abrechenbar']) else '',
                'Zeiterfasser': str(row['Zeiterfasser']) if not pd.isna(row['Zeiterfasser']) else '',
                'Kunde': str(row['Kunde']) if not pd.isna(row['Kunde']) else '',
                'Projekt': str(row['Projekt']) if not pd.isna(row['Projekt']) else '',
                'Aufgabe':  str(row['Aufgabe']) if not pd.isna(row['Aufgabe']) else '',
                'Dienstleistung': str(row['Dienstleistung']) if not pd.isna(row['Dienstleistung']) else '',
                'Kommentar': str(row['Kommentar']) if not pd.isna(row['Kommentar']) else ''
            }
        output_filename = "weekly_report.docx"

        create_weekly_report(weekly_data, output_filename)
        status_label.config(text=f"Successfully converted to {output_filename}")
    except FileNotFoundError:
        status_label.config(text="File not found. Please check the file path.")
    except KeyError as e:
        status_label.config(text=f"KeyError: {str(e)}. Please check CSV column names.")
    except Exception as e:
        status_label.config(text=f"An error occurred: {str(e)}")


# Tkinter GUI setup
root = tk.Tk()
root.title("CSV to Word Converter")

# File path selection
file_path_label = tk.Label(root, text="CSV File Path:")
file_path_label.pack()

file_path_entry = tk.Entry(root, width=50)
file_path_entry.pack()

browse_button = tk.Button(root, text="Browse", command=browse_file)
browse_button.pack()

# Convert button
convert_button = tk.Button(root, text="Convert to Word", command=convert_to_word)
convert_button.pack()

# Status label
status_label = tk.Label(root, text="")
status_label.pack()

root.mainloop()
